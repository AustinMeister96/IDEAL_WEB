from django.shortcuts import render
from django.shortcuts import HttpResponse
from .forms import ParticipantForm, QuestionForm, Breath_Collection_Form, Blood_Collection_Form, Lab_Processing_Form, Exposure_Form, Exposure_Form2, Exposure_Form3, Mandatory_questionaire_form, SearchForm, CT_Scan_Form, CancerHistoryForm, Mandatory_questionaire_form_c, Mandatory_questionaire_form_de, Mandatory_questionaire_form_fg, inclusion_form, testParticipantForm, MyForm, CT_Scan_Nodule_Form_1, CT_Scan_Nodule_Form_2
from .models import Participant, Question, Breath_Collection, Blood_Collection, lab_processing, Mandatory_questionaire, ct_scan, ct_scan_nodule_1, ct_scan_nodule_2, ct_scan_nodule_3, ct_scan_nodule_4, ct_scan_nodule_5, Exposure, Exposure2, Exposure3, Mandatory_questionaire_c, Mandatory_questionaire_de, Mandatory_questionaire_fg, inclusion,  testParticipant, UserAccounts, Exposure, Exposure2, Exposure3, Mandatory_questionaire, Mandatory_questionaire_c, Mandatory_questionaire_de, Mandatory_questionaire_fg, inclusion,  testParticipant, UserAccounts,PLCO_score
from django.shortcuts import render, get_object_or_404
from django.forms import modelformset_factory
from .forms import biological_relatives_with_cancer_form, plco_score_form, ParticipantSearchForm, CT_Scan_Nodule_Form_3, CT_Scan_Nodule_Form_4, CT_Scan_Nodule_Form_5
from django.core.exceptions import ObjectDoesNotExist

from django.http import HttpResponseRedirect as redirect
from django.http import HttpResponseRedirect
from django.urls import reverse
from django.http import JsonResponse
from django.contrib.sessions.models import Session
import re
from django.http import FileResponse
from pdfrw import PdfReader, PdfWriter, IndirectPdfDict
import json
from django.db.models import Q
from django.db import IntegrityError
from django.contrib import messages
from .models import History, biological_relatives_with_cancer
from django.views.generic.edit import CreateView, UpdateView
import logging
from django.contrib.auth.decorators import login_required
from django.contrib.sessions.models import Session
from django.template.loader import get_template
from django.views.generic import View
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from io import BytesIO
import json
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.pagesizes import letter
from reportlab.lib import styles
from django.contrib.staticfiles import finders
from django.http import JsonResponse
from xhtml2pdf import pisa
from django.template.loader import get_template

#def index(request):
 #   return HttpResponse("Hello, world. You're at the polls index.")

#create an index return that displays the html file called index.html


def download_pdf(request):
    # Find the absolute path to the PDF file
    pdf_path = finders.find('your_file.pdf')
    print('orking here 1')
    if pdf_path:
        with open(pdf_path, 'rb') as pdf_file:
            print('working here')
            response = HttpResponse(pdf_file.read(), content_type='application/pdf')
            response['Content-Disposition'] = 'attachment; filename="your_file.pdf"'
            return response
    else:
        return HttpResponse("PDF file not found.", status=404)


class GeneratePDF2(View):
    def get(self, request, *args, **kwargs):
        form_data = get_form_data(request)

        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="data.pdf"'

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        elements = []
        elements.append(Paragraph("Form Data PDF", styles['Title']))
        elements.append(Paragraph("Below is the data submitted in the form:", styles['Normal']))
        
        for field, value in form_data.items():
            field_line = f"<strong>{field}:</strong> {value}<br/>"
            elements.append(Paragraph(field_line, styles['Normal']))

        doc.build(elements)
        pdf = buffer.getvalue()
        buffer.close()

        response.write(pdf)
        return response


class GeneratePDF(View):
    def get(self, request, *args, **kwargs):
        form_data_json = request.session.get('form_data')
        if form_data_json:
            form_data = json.loads(form_data_json)
            form_data_string = "\n".join([f"{key}: {value}" for key, value in form_data.items()])

            response = HttpResponse(content_type='application/pdf')
            response['Content-Disposition'] = 'attachment; filename="data.pdf"'

            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()

            elements = []
            elements.append(Paragraph("Form Data PDF", styles['Title']))
            elements.append(Paragraph("Below is the data submitted in the form:", styles['Normal']))
            elements.append(Paragraph(form_data_string, styles['Normal']))  # Include the form data

            doc.build(elements)
            pdf = buffer.getvalue()
            buffer.close()

            response.write(pdf)
            return response
        else:
            return HttpResponse("Form data not available.")


def generate_pdf_oldNotWorking(data):
    print('working here')
    buffer = BytesIO()

    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []

    # Add form data to the PDF
    for field_name, value in data.items():
        if field_name != 'csrfmiddlewaretoken':  # Exclude CSRF token field
            field_label = field_name.replace('_', ' ').title()
            elements.append(f'{field_label}: {value}')

    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_pdf(data):
    buffer = BytesIO()

    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []

    # Define a style for the paragraphs
    paragraph_style = styles.getSampleStyleSheet()['Normal']

    # Add form data to the PDF
    for field_name, value in data.items():
        if field_name != 'csrfmiddlewaretoken':  # Exclude CSRF token field
            field_label = field_name.replace('_', ' ').title()

            # Create a paragraph for each piece of information
            paragraph_text = f'<b>{field_label}:</b> {value}'
            paragraph = Paragraph(paragraph_text, style=paragraph_style)
            elements.append(paragraph)

    doc.build(elements)
    buffer.seek(0)
    print('working pdf here')
    return buffer

def get_form_data(request):
    form_data = {
        'Visit Date': request.POST.get('{form.visit_date}', ''),
        'Postal Code': request.POST.get('postal_code', ''),
        'Date of Birth': request.POST.get('date_of_birth', ''),
        'Current Age': request.POST.get('current_age', ''),
        'Current Height': request.POST.get('current_height', ''),
        'Current Height Unit': request.POST.get('current_height_unit', ''),
        'Current Weight': request.POST.get('current_weight', ''),
        'Current Weight Unit': request.POST.get('current_weight_unit', ''),
        'Sex at Birth': request.POST.get('sex_birth', ''),
        'Gender Identity': request.POST.get('gender_identity', ''),
        'Ethnicity': request.POST.get('ethnicity', ''),
        'Ethnicity (Other)': request.POST.get('ethnicity_other', ''),
        'Born in Canada': request.POST.get('born_in_canada', ''),
        'Year Moved to Canada': request.POST.get('year_moved_to_canada', ''),
        'Birthplace': request.POST.get('birthplace', ''),
        'Highest Education Level': request.POST.get('highest_education_lvl', ''),
        'Highest Education Level (Other)': request.POST.get('highest_education_lvl_other', ''),
    }
    return form_data




def index(request):
    form = ParticipantForm
    return render(request, 'DataEntry/index.html', {'form': form})

def search_and_add(request):
    form = ParticipantSearchForm()

    if request.method == 'POST':
        form = ParticipantSearchForm(request.POST)
        if form.is_valid():
            participant_number = form.cleaned_data['participant_number']

            # Query the database for the participant number
            # Assuming you have a Participant model
            try:
                participant = Participant.objects.get(participant_number=participant_number)
                # Add participant number to session
                request.session['participant_num'] = participant.participant_number
                return render(request, 'search_result.html', {'participant': participant})
            except Participant.DoesNotExist:
                # Handle the case when the participant is not found
                pass

    return render(request, 'search_page.html', {'form': form})

def add_inclusion_old(request):
    if request.method == 'POST':
        participant_form = ParticipantForm(request.POST)
        inclusion_form_data = inclusion_form(request.POST)

        if participant_form.is_valid():
            participant_num = participant_form.cleaned_data['participant_number']
            try:
                participant_num = int(participant_num)  # Convert to integer)
                print('wroked here')
                request.session['participant_num'] = participant_num
            except ValueError:
                return JsonResponse({'error': 'Invalid participant number'}, status=400)

            participant, created = Participant.objects.get_or_create(participant_number=participant_num)

            if inclusion_form_data.is_valid():
                inclusion_data = inclusion_form_data.save(commit=False)
                inclusion_data.participant_number = participant  # Update the ForeignKey
                inclusion_data.save()

                return JsonResponse({'message': 'Inclusion form submitted successfully'})
            else:
                print("Inclusion Form Errors:", inclusion_form_data.errors)
                return JsonResponse({'error': inclusion_form_data.errors}, status=400)
        else:
            print("Participant Form Errors:", participant_form.errors)
            return JsonResponse({'error': participant_form.errors}, status=400)

    else:
        form = inclusion_form()
        participant_form = ParticipantForm()

    context = {
        'form': inclusion_form(),
        'participant_form': ParticipantForm(),
    }
    print(participant_form)
    return render(request, 'DataEntry/inclusion.html', context)


def add_inclusion_oldSept29(request, participant_num=None):
    participant_already_exists = False
    
    if request.method == 'POST':
        participant_form = ParticipantForm(request.POST)
        
        if participant_form.is_valid():
            participant_num = participant_form.cleaned_data['participant_number']
            cleaned_participant_num = participant_num.replace('-', '0')
            
            participant, created = Participant.objects.get_or_create(
                participant_number=cleaned_participant_num
                
            )
            
            if created:
                # Redirect to the same page but append participant number to URL
                return redirect(f'/DataEntry/add_inclusion/{cleaned_participant_num}/')
            else:
                participant_already_exists = True
                print("This participant number has already been used")
        else:
            print("Participant Form Errors:", participant_form.errors)
            return JsonResponse({'error': participant_form.errors}, status=400)

    else:
        participant_form = ParticipantForm()

    context = {
        'participant_form': participant_form,
        'participant_already_exists': participant_already_exists,
    }

    return render(request, 'DataEntry/inclusion.html', context)


def add_inclusion(request, participant_num=None):
    participant_already_exists = False

    if request.method == 'POST':
        participant_form = ParticipantForm(request.POST)

        if participant_form.is_valid():
            participant_num = participant_form.cleaned_data['participant_number']
            cleaned_participant_num = participant_num.replace('-', '0')

            participant, created = Participant.objects.get_or_create(
                participant_number=cleaned_participant_num
            )

            if created:
                # Redirect to the same page but append participant number to URL
                return redirect('add_inclusion', participant_num=cleaned_participant_num)
            else:
                participant_already_exists = True
                print("This participant number has already been used")
        else:
            print("Participant Form Errors:", participant_form.errors)

    else:
        participant_form = ParticipantForm()

    context = {
        'participant_form': participant_form,
        'participant_already_exists': participant_already_exists,
    }

    return render(request, 'DataEntry/inclusion.html', context)

def add_inclusion_workingNoIncData(request):
    participant_already_exists = False
    
    if request.method == 'POST':
        participant_form = ParticipantForm(request.POST)
        
        if participant_form.is_valid():
            participant_num = participant_form.cleaned_data['participant_number']
            print(participant_num)
            
            cleaned_participant_num = participant_num.replace('-', '0')
            print(cleaned_participant_num)
            
            participant, created = Participant.objects.get_or_create(
                participant_number=cleaned_participant_num
            )
            
            if not created:
                participant_already_exists = True
                print("This participant number has already been used")
        else:
            print("Participant Form Errors:", participant_form.errors)
            return JsonResponse({'error': participant_form.errors}, status=400)

    else:
        participant_form = ParticipantForm()
    
    context = {
        'participant_form': participant_form,
        'participant_already_exists': participant_already_exists,
    }
    
    return render(request, 'DataEntry/inclusion.html', context)



def add_inclusion_workingParticipantOnly(request):
    if request.method == 'POST':
        participant_form = ParticipantForm(request.POST)
        #inclusion_form_data = inclusion_form(request.POST)  # Assuming 'InclusionForm' is the correct form class

        if participant_form.is_valid():
            participant_num = participant_form.cleaned_data['participant_number']
            print(participant_num)

            # Remove hyphens and replace them with '0'
            cleaned_participant_num = participant_num.replace('-', '0')
            print(cleaned_participant_num)

            # Check if the Participant instance already exists, if not, create a new one.
            participant, created = Participant.objects.get_or_create(
                participant_number=cleaned_participant_num
            )

            # if inclusion_form_data.is_valid():
            #     inclusion_data = inclusion_form_data.save(commit=False)
                
            #     # Assume cleaned_participant_num is already prepared, e.g., '1004856'
            #     participant, created = Participant.objects.get_or_create(participant_number=cleaned_participant_num)
                
            #     inclusion_data.participant_number = participant  # Link the Participant object
            #     inclusion_data.save()
                
            #     return JsonResponse({'message': 'Inclusion form submitted successfully'})
            # else:
            #     print("Inclusion Form Errors:", inclusion_form_data.errors)
            #     return JsonResponse({'error': inclusion_form_data.errors}, status=400)
        else:
            print("Participant Form Errors:", participant_form.errors)
            return JsonResponse({'error': participant_form.errors}, status=400)

    else:
        #inclusion_form_instance = inclusion_form()  # Assuming 'InclusionForm' is the correct form class
        participant_form = ParticipantForm()

    context = {
        #'form': inclusion_form_instance,
        'participant_form': participant_form,
    }
    return render(request, 'DataEntry/inclusion.html', context)




def generate_word_document(data):
    doc = Document()
    doc.add_heading('Mandatory Questionnaire', level=1)
    doc.add_paragraph(f"Name: {data['name']}")
    doc.add_paragraph(f"Email: {data['email']}")
    # Add more paragraphs for other form fields
    return doc


def add_mandatory_questionaire(request, participant_id=None):
    print(request.GET)  # Print the GET parameters
    
    participant = None
    mandatory_questionaire_data = None
    print("View accessed!")
    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)
        mandatory_questionaire_data = Mandatory_questionaire.objects.filter(participant_num=participant).first()
        print(participant)
    RelativeCancerFormSet = modelformset_factory(biological_relatives_with_cancer, form=biological_relatives_with_cancer_form, extra=1)
    
    if request.method == 'POST':
        form = Mandatory_questionaire_form(request.POST, instance=mandatory_questionaire_data)
        formset = RelativeCancerFormSet(request.POST, prefix='relatives')

        if form.is_valid() and formset.is_valid():
            print('form is valid')
            if not participant:
                participant = Participant.objects.create(participant_number=participant_id)
            
            mandatory_questionaire = form.save(commit=False)
            mandatory_questionaire.participant_num = participant
            mandatory_questionaire.save()

            instances = formset.save(commit=False)
            for instance in instances:
                instance.mandatory_questionaire = mandatory_questionaire
                instance.save()

            word_doc = generate_word_document(form.cleaned_data)

            # Serve the Word document for download
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename=mandatory_questionnaire.docx'
            word_doc.save(response)

            return HttpResponseRedirect(reverse('add_mandatory_questionaire', args=[participant_id]) + '?submitted=True')
        else:
            print(form.errors)
            print(formset.errors)

    elif 'generate_pdf' in request.GET:  # Handle the PDF generation
        print('button pressed')
        form = Mandatory_questionaire_form(instance=mandatory_questionaire_data)  # Create a form instance with the pdf_data
        if mandatory_questionaire_data:
            formset = RelativeCancerFormSet(queryset=biological_relatives_with_cancer.objects.filter(mandatory_questionaire=mandatory_questionaire_data), prefix='relatives')
        else:
            formset = RelativeCancerFormSet(queryset=biological_relatives_with_cancer.objects.none(), prefix='relatives')
         # Validate the form
        pdf_data = form.instance 
        pdf_buffer = generate_mandatory_questionnaire_pdf(pdf_data, request)
        print(pdf_buffer)
        response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
        print(response)
        response['Content-Disposition'] = f'attachment; filename=mandatory_questionnaire.pdf'
        print('should be wokring here')
        return response

    else:
        print("Form is NOT valid!")
        form = Mandatory_questionaire_form(instance=mandatory_questionaire_data)
        if mandatory_questionaire_data:
            formset = RelativeCancerFormSet(queryset=biological_relatives_with_cancer.objects.filter(mandatory_questionaire=mandatory_questionaire_data), prefix='relatives')
        else:
            formset = RelativeCancerFormSet(queryset=biological_relatives_with_cancer.objects.none(), prefix='relatives')

    context = {
        'form': form,
        'formset': formset,
        'participant_id': participant_id,
    }
    print(form.errors)
    print(formset.errors)
    print(request.POST)
    return render(request, 'DataEntry/mandatory_questionaire.html', context)


def generate_mandatory_questionnaire_pdf(pdf_data, request):
    template_path = 'DataEntry/pdf_template.html'  # Replace with the actual path
    template = get_template(template_path)
    context = {'form': pdf_data}  # Pass the form instance to the template context
    html = template.render(context)
    pdf_buffer = BytesIO()
    # Create a PDF from the HTML content
    pisa_status = pisa.CreatePDF(html, dest=pdf_buffer)
    if pisa_status.err:
        return HttpResponse('PDF generation error', content_type='text/plain')
    response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename=mandatory_questionnaire.pdf'
    pdf_buffer.close()
    return response



def add_mandatory_questionaire_workingNowAug29(request, participant_id=None):
    print(request.POST)
    
    participant = None
    mandatory_questionaire_data = None
    print("View accessed!")
    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)
        mandatory_questionaire_data = Mandatory_questionaire.objects.filter(participant_num=participant).first()
        print(participant)
    RelativeCancerFormSet = modelformset_factory(biological_relatives_with_cancer, form=biological_relatives_with_cancer_form, extra=1)
    
    if request.method == 'POST':
        form = Mandatory_questionaire_form(request.POST, instance=mandatory_questionaire_data)
        formset = RelativeCancerFormSet(request.POST, prefix='relatives')

        if form.is_valid() and formset.is_valid():
            if 'generate_pdf' in request.POST:
                    pdf_data = form.cleaned_data  # Use cleaned_data here after form validation
                    print("Form data:", pdf_data)
                    pdf_buffer = generate_pdf(pdf_data)
                    print("PDF data:", pdf_data)

                    # Serve the PDF document for download
                    response = HttpResponse(pdf_buffer.read(), content_type='application/pdf')
                    response['Content-Disposition'] = f'attachment; filename=mandatory_questionnaire.pdf'
                    print("PDF buffer size:", pdf_buffer.tell())
                    return response
            if not participant:
                participant = Participant.objects.create(participant_number=participant_id)
            
            mandatory_questionaire = form.save(commit=False)
            mandatory_questionaire.participant_num = participant
            mandatory_questionaire.save()

            instances = formset.save(commit=False)
            for instance in instances:
                instance.mandatory_questionaire = mandatory_questionaire
                instance.save()

            word_doc = generate_word_document(form.cleaned_data)

            # Serve the Word document for download
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename=mandatory_questionnaire.docx'
            word_doc.save(response)

            return HttpResponseRedirect(reverse('add_mandatory_questionaire', args=[participant_id]) + '?submitted=True')
        else:
            print(form.errors)
            print(formset.errors)
            print(request.POST)
    else:
        print("Form is NOT valid!")
        print(request.POST)
        form = Mandatory_questionaire_form(instance=mandatory_questionaire_data)
        if mandatory_questionaire_data:
            formset = RelativeCancerFormSet(queryset=biological_relatives_with_cancer.objects.filter(mandatory_questionaire=mandatory_questionaire_data), prefix='relatives')
        else:
            formset = RelativeCancerFormSet(queryset=biological_relatives_with_cancer.objects.none(), prefix='relatives')
        #print(form.errors)



    context = {
        'form': form,
        'formset': formset,
        'participant_id': participant_id,
    }

    #print(form.errors)
    #print(formset.errors)
    print("Form errors:", form.errors)
    print("Formset errors:", formset.errors)
    
    return render(request, 'DataEntry/mandatory_questionaire.html', context)

def add_mandatory_questionaire_c_original(request, participant_id=None):
    print(f"URL Parameter - participant_id: {participant_id}")
    print(f"Participant ID: {participant_id}")
    print(f"Request method: {request.method}")
    
    participant = None
    mandatory_questionaire_c_data = None
    print('working here 1')
    
    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)
        mandatory_questionaire_c_data = Mandatory_questionaire_c.objects.filter(participant_num=participant).first()
        print(participant)
    
    if request.method == 'POST':
        form = Mandatory_questionaire_form_c(request.POST, instance=mandatory_questionaire_c_data)
        if not participant:
            participant = Participant.objects.create(participant_number=participant_id)
            print(f"Participant object: {participant}")
        
        if form.is_valid():
            mandatory_questionaire_c = form.save(commit=False)
            mandatory_questionaire_c.participant_num = participant
            mandatory_questionaire_c.save()
            print(f"Questionnaire's participant number: {mandatory_questionaire_c.participant_num}")
        else:
            print(form.errors)
    else:
        form = Mandatory_questionaire_form_c(instance=mandatory_questionaire_c_data)
    
    context = {
        'form': form,
        'participant_id': participant_id,
    }
    print(form.errors)
    return render(request, 'DataEntry/mandatory_questionaire_c.html', context)

def add_mandatory_questionaire_c_old2(request, participant_id=None):
    print(f"URL Parameter - participant_id: {participant_id}")
    print(f"Participant ID: {participant_id}")
    print(f"Request method: {request.method}")
    
    participant = None
    mandatory_questionaire_c_data = None
    print('working here 1')
    
    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)
        mandatory_questionaire_c_data = Mandatory_questionaire_c.objects.filter(participant_num=participant).first()
        print(participant)
    
    if request.method == 'POST':
        form = Mandatory_questionaire_form_c(request.POST, instance=mandatory_questionaire_c_data)
        if not participant:
            participant = Participant.objects.create(participant_number=participant_id)
            print(f"Participant object: {participant}")
        
        if form.is_valid():
            mandatory_questionaire_c = form.save(commit=False)
            mandatory_questionaire_c.participant_num = participant
            mandatory_questionaire_c.save()
            print(f"Questionnaire's participant number: {mandatory_questionaire_c.participant_num}")
        else:
            print(form.errors)
    elif 'generate_pdf' in request.GET:  # Handle the PDF generation
        print('button pressed')
        form = Mandatory_questionaire_c_form(instance=mandatory_questionaire_c_data)  # Create a form instance with the pdf_data
        if mandatory_questionaire_c_data:
            formset = RelativeCancerFormSet(queryset=biological_relatives_with_cancer.objects.filter(mandatory_questionaire=mandatory_questionaire_c_data), prefix='relatives')
        else:
            formset = RelativeCancerFormSet(queryset=biological_relatives_with_cancer.objects.none(), prefix='relatives')
            # Validate the form
        pdf_data = form.instance
        pdf_buffer = generate_mandatory_questionnaire_c_pdf(pdf_data, request)
        print(pdf_buffer)
        response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
        print(response)
        response['Content-Disposition'] = f'attachment; filename=mandatory_questionnaire.pdf'
        print('should be wokring here')
        return response
    else:
        form = Mandatory_questionaire_form_c(instance=mandatory_questionaire_c_data)
    
    context = {
        'form': form,
        'participant_id': participant_id,
    }
    print(form.errors)
    return render(request, 'DataEntry/mandatory_questionaire_c.html', context)

def add_mandatory_questionaire_c(request, participant_id=None):
    print(request.GET)  # Print the GET parameters
    
    participant = None
    mandatory_questionaire_c_data = None
    print("View accessed!")
    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)
        mandatory_questionaire_c_data = Mandatory_questionaire_c.objects.filter(participant_num=participant).first()
        print(participant)
    
    
    if request.method == 'POST':
        form = Mandatory_questionaire_form_c(request.POST, instance=mandatory_questionaire_c_data)
        if not participant:
            participant = Participant.objects.create(participant_number=participant_id)
            print('created participant')

        if form.is_valid():
            print('form is valid')
            if not participant:
                participant = Participant.objects.create(participant_number=participant_id)
            
            mandatory_questionaire_c = form.save(commit=False)
            mandatory_questionaire_c.participant_num = participant
            mandatory_questionaire_c.save()
        else:
            print(form.errors)
            
    elif 'generate_pdf' in request.GET:  # Handle the PDF generation
        print('button pressed')
        form = Mandatory_questionaire_form_c(instance=mandatory_questionaire_c_data)  # Create a form instance with the pdf_data

        pdf_data = form.instance
        pdf_buffer = generate_mandatory_questionnaire_c_pdf(pdf_data, request)
        response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename=mandatory_questionnaire.pdf'
        return response

    else:
        print("Form is NOT valid!")
        form = Mandatory_questionaire_form_c(instance=mandatory_questionaire_c_data)

    context = {
        'form': form,
        'participant_id': participant_id,
    }

    return render(request, 'DataEntry/mandatory_questionaire_c.html', context)


def add_mandatory_questionaire_c_with_id(request, participant_id):
    return add_mandatory_questionaire_c(request, participant_id)


def generate_mandatory_questionnaire_c_pdf(pdf_data, request):
    template_path = 'DataEntry/mandatory_questionaire_c_pdf.html'  # Replace with the actual path
    template = get_template(template_path)
    context = {'form': pdf_data}  # Pass the form instance to the template context
    html = template.render(context)
    pdf_buffer = BytesIO()
    # Create a PDF from the HTML content
    pisa_status = pisa.CreatePDF(html, dest=pdf_buffer)
    if pisa_status.err:
        return HttpResponse('PDF generation error', content_type='text/plain')
    response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename=mandatory_questionnaire.pdf'
    pdf_buffer.close()
    return response













def add_mandatory_questionaire_de(request, participant_id=None):
    print(request.POST)
    print(f"URL Parameter - participant_id: {participant_id}")
    print(f"Participant ID: {participant_id}")
    print(f"Request method: {request.method}")
    participant = None
    mandatory_questionaire_de_data = None
    mandatory_questionaire_fg_data = None
    print("View accessed!")

    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)
        mandatory_questionaire_de_data = Mandatory_questionaire_de.objects.filter(participant_num=participant).first()
        mandatory_questionaire_fg_data = Mandatory_questionaire_fg.objects.filter(participant_num=participant).first()
        print(participant)

    if request.method == 'POST':
        form = Mandatory_questionaire_form_de(request.POST, instance=mandatory_questionaire_de_data, initial={'participant_num': participant})
        fg_form = Mandatory_questionaire_form_fg(request.POST, instance=mandatory_questionaire_fg_data, initial={'participant_num': participant})

        if form.is_valid() and fg_form.is_valid():
            # Save forms
            form.save()
            fg_form.save()

            if participant_id:
                return HttpResponseRedirect(reverse('add_mandatory_questionaire_de', args=[participant_id]) + '?submitted=True')
            else:
                return HttpResponseRedirect(reverse('add_mandatory_questionaire_de') + '?submitted=True')

        else:
            print("Form DE errors:", form.errors)
            print("Form FG errors:", fg_form.errors)

    elif 'generate_pdf' in request.GET:
        form = Mandatory_questionaire_form_de(instance=mandatory_questionaire_de_data)
        fg_form = Mandatory_questionaire_form_fg(instance=mandatory_questionaire_fg_data)
        pdf_buffer = generate_mandatory_questionnaire_d_pdf(form, fg_form, request)
        response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename=mandatory_questionnaire.pdf'
        return response

    else:
        print("Form is NOT valid!")
        form = Mandatory_questionaire_form_de(instance=mandatory_questionaire_de_data, initial={'participant_num': participant})
        fg_form = Mandatory_questionaire_form_fg(instance=mandatory_questionaire_fg_data, initial={'participant_num': participant})

    context = {
        'form': form,
        'fg_form': fg_form,
        'participant_id': participant_id,
    }
    print(form.errors)
    print(fg_form.errors)
    return render(request, 'DataEntry/mandatory_questionaire_de.html', context)


def add_exposure_form_old(request, participant_id=None):
    form = Exposure_Form(request.POST or None)
    form2 = Exposure_Form2(request.POST or None)
    form3 = Exposure_Form3(request.POST or None)

    if request.method == 'POST':
        if form.is_valid() and form2.is_valid() and form3.is_valid():
            form.save()
            return HttpResponse('/exposure_form/?submitted=True')
    context = {
        'form1': form,
        'form': form2,
        'form2': form3,
    }
    return render(request, 'DataEntry/exposure_form.html', context) #{'form': form}

# def generate_mandatory_questionnaire_d_pdf(pdf_data, request):
#     template_path = 'DataEntry/mandatory_questionaire_d_pdf.html'  # Replace with the actual path
#     template = get_template(template_path)
#     context = {'form': pdf_data}  # Pass the form instance to the template context
#     html = template.render(context)
#     pdf_buffer = BytesIO()
#     # Create a PDF from the HTML content
#     pisa_status = pisa.CreatePDF(html, dest=pdf_buffer)
#     if pisa_status.err:
#         return HttpResponse('PDF generation error', content_type='text/plain')
#     response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
#     response['Content-Disposition'] = 'attachment; filename=mandatory_questionnaire_d.pdf'
#     pdf_buffer.close()
#     return response

def generate_mandatory_questionnaire_d_pdf(form, fg_form, request):
    template_path = 'DataEntry/mandatory_questionaire_d_pdf.html'  # Replace with the actual path
    template = get_template(template_path)
    context = {
        'form': form,
        'fg_form': fg_form,
    }  # Pass both form instances to the template context
    html = template.render(context)
    pdf_buffer = BytesIO()
    # Create a PDF from the HTML content
    pisa_status = pisa.CreatePDF(html, dest=pdf_buffer)
    if pisa_status.err:
        return HttpResponse('PDF generation error', content_type='text/plain')
    response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename=mandatory_questionnaire_d.pdf'
    pdf_buffer.close()
    return response






def add_exposure_form(request, participant_id=None):
    participant = None
    exposure2_data = None
    exposure3_data = None

    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)
        exposure2_data = Exposure2.objects.filter(participant_num=participant).first()
        exposure3_data = Exposure3.objects.filter(participant_num=participant).first()

    if request.method == 'POST':
        form2 = Exposure_Form2(request.POST, instance=exposure2_data, initial={'participant_num': participant})
        form3 = Exposure_Form3(request.POST, instance=exposure3_data, initial={'participant_num': participant})

        if form2.is_valid() and form3.is_valid():
            exposure2_instance = form2.save(commit=False)
            exposure2_instance.participant_num = participant
            exposure2_instance.save()

            exposure3_instance = form3.save(commit=False)
            exposure3_instance.participant_num = participant
            exposure3_instance.save()

            if participant_id:
                return HttpResponseRedirect(reverse('add_exposure_form', args=[participant_id]) + '?submitted=True')
            else:
                return HttpResponseRedirect(reverse('add_exposure_form') + '?submitted=True')
        else:
            print("Form2 errors:", form2.errors)
            print("Form3 errors:", form3.errors)
    elif 'generate_pdf' in request.GET:
        print('button pressed')
        form2 = Exposure_Form2(instance=exposure2_data)
        form3 = Exposure_Form3(instance=exposure3_data)
        pdf_buffer = generate_exposure_questionnaire_pdf(form2, form3, request)
        print(pdf_buffer)
        response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
        print(response)
        response['Content-Disposition'] = f'attachment; filename=mandatory_questionnaire.pdf'
        print('should be working here')
        return response

    
    else:
        form2 = Exposure_Form2(instance=exposure2_data, initial={'participant_num': participant})
        form3 = Exposure_Form3(instance=exposure3_data, initial={'participant_num': participant})

    context = {
        'form2': form2,
        'form3': form3,
        'participant_id': participant_id,
    }
  
    return render(request, 'DataEntry/exposure_form.html', context)

def generate_exposure_questionnaire_pdf(form2, form3, request):
    template_path = 'DataEntry/exposure_questionaire_pdf.html'  # Replace with the actual path
    template = get_template(template_path)
    context = {
        'form2': form2,
        'form3': form3,
    }  # Pass both form instances to the template context
    html = template.render(context)
    pdf_buffer = BytesIO()
    # Create a PDF from the HTML content
    pisa_status = pisa.CreatePDF(html, dest=pdf_buffer)
    if pisa_status.err:
        return HttpResponse('PDF generation error', content_type='text/plain')
    response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename=mandatory_questionnaire_d.pdf'
    pdf_buffer.close()
    return response


def add_exposure_form_oldWorking(request, participant_id=None):
    print(request.POST)
    print(f"URL Parameter - participant_id: {participant_id}")
    print(f"Participant ID: {participant_id}")
    print(f"Request method: {request.method}")
    participant = None
    exposure_data = None
    exposure2_data = None
    exposure3_data = None


    print("View accessed!")

    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)
        
        exposure2_data = Exposure2.objects.filter(participant_num=participant).first()
        exposure3_data = Exposure3.objects.filter(participant_num=participant).first()
        print(participant)

    if request.method == 'POST':
        
        form2 = Exposure_Form2(request.POST, instance=exposure2_data, initial={'participant_num': participant})
        form3 = Exposure_Form3(request.POST, instance=exposure3_data, initial={'participant_num': participant})

        if form2.is_valid() and form3.is_valid():

            exposure2_instance = form2.save(commit=False)
            exposure2_instance.participant_num = participant
            exposure2_instance.save()

            exposure3_instance = form3.save(commit=False)
            exposure3_instance.participant_num = participant
            exposure3_instance.save()

            if participant_id:
                return HttpResponseRedirect(reverse('add_exposure_form', args=[participant_id]) + '?submitted=True')
            else:
                return HttpResponseRedirect(reverse('add_exposure_form') + '?submitted=True')

        else:
            print("Form2 errors:", form2.errors)
            print("Form3 errors:", form3.errors)

    else:
        print("Form is NOT valid!")
        
        form2 = Exposure_Form2(request.POST or None)
        form3 = Exposure_Form3(request.POST or None)

    context = {
        'form2': form2,
        'form3': form3,
        'participant_id': participant_id,
    }
  
    print(form2.errors)
    print(form3.errors)
    return render(request, 'DataEntry/exposure_form.html', context)




def add_yearly_update(request):
    return render(request, 'DataEntry/yearly_update.html', {})

def add_breath_collection_old(request):
    form = Breath_Collection_Form
    if request.method == 'POST':
        form_class = Breath_Collection_Form
        form = Breath_Collection_Form(request.POST)
        
        if form.is_valid():
            print('test')
            form.save(commit=True)
            form.save()
            return HttpResponse('/add_breath_collection/?submitted=True')
    return render(request, 'DataEntry/breath_collection.html', {'form': form})

def add_breath_collection(request, participant_id=None):
    participant = None
    breath_collection_data = None

    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)
        breath_collection_data = Breath_Collection.objects.filter(participant_num=participant).first()

    form = Breath_Collection_Form(instance=breath_collection_data, initial={'participant_num': participant})

    if request.method == 'POST':
        form = Breath_Collection_Form(request.POST, instance=breath_collection_data, initial={'participant_num': participant})

        if form.is_valid():
            form.save()

            if participant_id:
                return HttpResponseRedirect(reverse('add_breath_collection', args=[participant_id]) + '?submitted=True')
            else:
                return HttpResponseRedirect(reverse('add_breath_collection') + '?submitted=True')

    context = {
        'form': form,
        'participant_id': participant_id,
    }
    print(form.errors)
    print(participant_id)

    return render(request, 'DataEntry/breath_collection.html', context)

def add_ct_scan(request, participant_id=None):
    print("participant_id received:", participant_id)
    participant = get_object_or_404(Participant, participant_number=participant_id)
    print("Participant object retrieved:", participant) # Debugging Statement
    participant = get_object_or_404(Participant, participant_number=participant_id)
    ct_scan_data = ct_scan.objects.filter(participant_num=participant).first()
    ct_nodule_form_1 = ct_scan_nodule_1.objects.all()
    ct_nodule_form_2 = ct_scan_nodule_2.objects.all()
    ct_nodule_form_3 = ct_scan_nodule_3.objects.all()
    ct_nodule_form_4 = ct_scan_nodule_4.objects.all()
    ct_nodule_form_5 = ct_scan_nodule_5.objects.all()

    if request.method == 'POST':
        ct_scan_form = CT_Scan_Form(request.POST, instance=ct_scan_data)
        ct_nodule_form = CT_Scan_Nodule_Form_1(request.POST)
        ct_nodule_form2 = CT_Scan_Nodule_Form_2(request.POST)
        ct_nodule_form3 = CT_Scan_Nodule_Form_3(request.POST)
        ct_nodule_form4 = CT_Scan_Nodule_Form_4(request.POST)
        ct_nodule_form5 = CT_Scan_Nodule_Form_5(request.POST)

        if 'save_ct_scan' in request.POST:
            if ct_scan_form.is_valid():
                ct_scan_data = ct_scan_form.save(commit=False)
                ct_scan_data.participant_num_id = participant.pk
                ct_scan_data.save()
                messages.success(request, "CT Scan Form saved successfully.")
                ct_scan_form = CT_Scan_Form(instance=ct_scan_form.instance)
                #return redirect('success_page')

        elif 'save_nodule' in request.POST:
            if ct_nodule_form.is_valid():
                ct_nodule_form.instance.participant_num_id = participant.pk
                ct_nodule_form2.instance.participant_num_id = participant.pk
                ct_nodule_form3.instance.participant_num_id = participant.pk
                ct_nodule_form4.instance.participant_num_id = participant.pk
                ct_nodule_form5.instance.participant_num_id = participant.pk
                ct_nodule_form.save()
                ct_nodule_form2.save()
                ct_nodule_form3.save()
                ct_nodule_form4.save()
                ct_nodule_form5.save()
                messages.success(request, "CT Nodule Form saved successfully.")
                ct_nodule_form = CT_Scan_Nodule_Form_1(instance=ct_nodule_form.instance)
                ct_nodule_form2 = CT_Scan_Nodule_Form_2(instance=ct_nodule_form.instance)
                ct_nodule_form3 = CT_Scan_Nodule_Form_3(instance=ct_nodule_form.instance)
                ct_nodule_form4 = CT_Scan_Nodule_Form_4(instance=ct_nodule_form.instance)
                ct_nodule_form5 = CT_Scan_Nodule_Form_5(instance=ct_nodule_form.instance)
                #return redirect('success_page')

    else:
        if ct_scan_data:
            ct_scan_form = CT_Scan_Form(instance=ct_scan_data)
            ct_nodule_form = CT_Scan_Nodule_Form_1(instance=ct_nodule_form.instance)
            ct_nodule_form2 = CT_Scan_Nodule_Form_2(instance=ct_nodule_form.instance)
            ct_nodule_form3 = CT_Scan_Nodule_Form_3(instance=ct_nodule_form.instance)
            ct_nodule_form4 = CT_Scan_Nodule_Form_4(instance=ct_nodule_form.instance)
            ct_nodule_form5 = CT_Scan_Nodule_Form_5(instance=ct_nodule_form.instance)
        else:
            ct_scan_form = CT_Scan_Form(initial={'participant_num': participant})
            ct_nodule_data = ct_scan_nodule_1.objects.filter(participant_num=participant).first()
            ct_nodule_data2 = ct_scan_nodule_2.objects.filter(participant_num=participant).first()
            ct_nodule_data3 = ct_scan_nodule_3.objects.filter(participant_num=participant).first()
            ct_nodule_data4 = ct_scan_nodule_4.objects.filter(participant_num=participant).first()
            ct_nodule_data5 = ct_scan_nodule_5.objects.filter(participant_num=participant).first()
        if ct_nodule_data:
            ct_nodule_form = CT_Scan_Nodule_Form_1(instance=ct_nodule_data)
            ct_nodule_form2 = CT_Scan_Nodule_Form_2(instance=ct_nodule_data2)
            ct_nodule_form3 = CT_Scan_Nodule_Form_3(instance=ct_nodule_data3)
            ct_nodule_form4 = CT_Scan_Nodule_Form_4(instance=ct_nodule_data4)
            ct_nodule_form5 = CT_Scan_Nodule_Form_5(instance=ct_nodule_data5)
        else:
            ct_nodule_form = CT_Scan_Nodule_Form_1(initial={'participant_num': participant})
            ct_nodule_form2 = CT_Scan_Nodule_Form_2(initial={'participant_num': participant})
            ct_nodule_form3 = CT_Scan_Nodule_Form_3(initial={'participant_num': participant})
            ct_nodule_form4 = CT_Scan_Nodule_Form_4(initial={'participant_num': participant})
            ct_nodule_form5 = CT_Scan_Nodule_Form_5(initial={'participant_num': participant})


    context = {
        'ct_scan_form': ct_scan_form,
        'participant': participant,
        'ct_nodule_form': ct_nodule_form,
        'ct_nodule_form2': ct_nodule_form2,
        'ct_nodule_form3': ct_nodule_form3,
        'ct_nodule_form4': ct_nodule_form4,
        'ct_nodule_form5': ct_nodule_form5,
    }
    print("participant:", participant)
    print("CT Scan Form Errors:", ct_scan_form.errors)
    print("CT Nodule Form Errors:", ct_nodule_form.errors)
    return render(request, 'DataEntry/ct_scan.html', context)


def add_blood_collection(request):
    form = Blood_Collection_Form
    if request.method == 'POST':
        form_class = Blood_Collection_Form
        form = Blood_Collection_Form(request.POST)
        
        if form.is_valid():
            print('test')
            return HttpResponse('/add_blood_collection/?submitted=True')
    return render(request, 'DataEntry/blood_collection.html', {'form': form})



def add_plco_score(request):
    if request.method == 'POST':
        # If the request method is POST, it means the form has been submitted
        form = plco_score_form(request.POST)
        if form.is_valid():
            # Save the form data to the database
            form.save()
            # Redirect to a success page or any other page you prefer
            return redirect('success_page')  # Replace 'success_page' with the URL name of the success page
    else:
        # If the request method is not POST, it means it's a GET request
        # Show the form to the user for data entry
        form = plco_score_form()

    return render(request, 'DataEntry/plco_score.html', {'form': form})

#need to test this:
# from django.shortcuts import render, redirect
# from .forms import plco_score_form  # Import your form class
# from .your_plco_calculator_module import plco_score  # Import your plco_score function

# def add_plco_score(request):
#     if request.method == 'POST':
#         If the request method is POST, it means the form has been submitted
#         form = plco_score_form(request.POST)
#         if form.is_valid():
#             Extract form data
#             age = form.cleaned_data['age']
#             education = form.cleaned_data['education']
#             bmi = form.cleaned_data['bmi']
#             copd = form.cleaned_data['copd']
#             personal_hx_cancer = form.cleaned_data['personal_hx_cancer']
#             family_hx_cancer = form.cleaned_data['family_hx_cancer']
#             race = form.cleaned_data['race']
#             smoke_status = form.cleaned_data['smoke_status']
#             avg_num_cigs = form.cleaned_data['avg_num_cigs']
#             duration = form.cleaned_data['duration']
#             yrs_quit = form.cleaned_data['yrs_quit']
            
#             Call your plco_score function
#             calculated_plco_score = plco_score(age, education, bmi, copd, personal_hx_cancer, family_hx_cancer, race, smoke_status, avg_num_cigs, duration, yrs_quit)
            
#             Save the form data and calculated PLCO score to the database
#             instance = form.save(commit=False)
#             instance.plco_score = calculated_plco_score  # Assuming 'plco_score' is the field name in your model
#             instance.save()
            
#             Redirect to a success page or any other page you prefer
#             return redirect('success_page')  # Replace 'success_page' with the URL name of the success page
#     else:
#         If the request method is not POST, it means it's a GET request
#         Show the form to the user for data entry
#         form = plco_score_form()

#     return render(request, 'DataEntry/plco_score.html', {'form': form})




def add_participant(request):
    submitted = False
    if request.method == 'POST':
        form_class = ParticipantForm
        form = ParticipantForm(request.POST)
        if form.is_valid():
            form.save()
            return HttpResponse('/add_participant/?submitted=True?id=' + str(form.instance.id))
    form = ParticipantForm
    context = {
        'form': form,
        'submitted': submitted

               }

    return render(request, 'DataEntry/index.html', {'form': form})


def add_protocol_deviation(request):
    submitted = False
    if request.method == 'POST':
        form_class = ParticipantForm
        form = ParticipantForm(request.POST)
        if form.is_valid():
            form.save()
            return HttpResponse('/add_participant/?submitted=True')
    form = ParticipantForm
    return render(request, 'DataEntry/index.html', {'form': form})







def update_participant_status(request):
    if request.method == 'POST':
        inclusion_criteria_values = [
            request.POST.get('inclusion_criteria_1'),
            request.POST.get('inclusion_criteria_2'),
            # Get values for inclusion_criteria_3 to inclusion_criteria_9_d
        ]

        # Evaluate the inclusion criteria conditions
        if all(value == 'yes' for value in inclusion_criteria_values[:5]) and all(value == 'no' for value in inclusion_criteria_values[5:]):
            participant_status = 'group 1'
        else:
            participant_status = 'group 2'

        # Update the participant_status in the database
        inclusion_object = inclusion.objects.first()  # Retrieve the inclusion object (You might need to modify this depending on your use case)
        inclusion_object.participant_status = participant_status
        inclusion_object.save()

        # Return updated participant_status as JSON response
        response_data = {'participant_status': participant_status}
        return JsonResponse(response_data)
    else:
        return render(request, 'form.html')



def add_lab_processing(request, participant_id=None):
    form = Lab_Processing_Form()

    if request.method == 'POST':
        form = Lab_Processing_Form(request.POST)

        if form.is_valid():
            participant_num = form.cleaned_data['participant_num']

            participant = get_object_or_404(lab_processing, participant_num=participant_num)
            if participant:
                form = Lab_Processing_Form(instance=participant)

                return render(request, 'add_lab_processing.html', {'form': form})
            else:
                form.save()
                return HttpResponse('/add_lab_processing/?submitted=True')
        else:
            return render(request, 'RECIVA.html', {'form': form})
    
    return render(request, 'DataEntry/lab_processing.html', {'form': form})




def update_mandatory_questionaire(request, participant_id):
    mandatory_questionaire = Mandatory_questionaire.objects.get(pk=participant_id)
    form = Mandatory_questionaire_form(request.POST or None, instance=mandatory_questionaire)

    return render(request, 'DataEntry/update_mandatory_questionaire.html', {'Mandatory_questionaire': Mandatory_questionaire, 'participant_id': participant_id})


def add_clinical_procedures(request):

    return render(request, 'DataEntry/clinical_procedures.html', {})


def add_protocol_deviations(request):
   

    return render(request, 'DataEntry/protocol_deviations.html', {})


def add_plco_score(request):
    if request.method == 'POST':
        # If the request method is POST, it means the form has been submitted
        form = plco_score_form(request.POST)
        if form.is_valid():
            # Save the form data to the database
            form.save()
            # Redirect to a success page or any other page you prefer
            return redirect('success_page')  # Replace 'success_page' with the URL name of the success page
    else:
        # If the request method is not POST, it means it's a GET request
        # Show the form to the user for data entry
        form = plco_score_form()

    return render(request, 'DataEntry/plco_score.html', {'form': form})







def add_lab_processing_with_data(request, participant_num):
    participant = lab_processing.objects.get(participant_num=participant_num)
    form = Lab_Processing_Form(instance=participant)

    return render(request, 'lab_processing_with_data.html', {'form': form})







def add_ct_scan_old(request):
    
    form = CT_Scan_Form(request.POST or None)
    form2 = CT_Scan_Nodule_Form(request.POST or None)
    if request.method == 'POST':
        if form.is_valid() and form2.is_valid():
            form.save()
            form2.save()
            return HttpResponse('/ct_scan/?submitted=True')
    context = {
        'form': form,
        'form2': form2,
    }
    return render(request, 'DataEntry/ct_scan.html', context) #{'form': form}


def add_participant(request):
    if request.method == 'POST':
        form = ParticipantForm(request.POST)
        if form.is_valid():
            form.save()
            return JsonResponse({'message': 'Participant form submitted successfully'})
        else:
            return JsonResponse({'error': form.errors}, status=400)




def search(request):
    query = request.GET.get('q')
    if query:
        participant = get_object_or_404(Participant, participant_number=query)
        return redirect('participant_detail', participant_id=participant.participant_number)
    else:
        return render(request, 'search.html')

def list_participants(request):
    participants = Participant.objects.all()
    context = {
        'participants': participants,
    }
    return render(request,'DataEntry/list_participants.html', context)

def show_participants(request, participant_id):
    participants = Participant.objects.get(pk=participant_id)
    context = {
        'participants': participants,
    }
    return render(request,'DataEntry/show_participants.html', context)


def history_page(request):
    histories = History.objects.all().order_by('-timestamp')
    return render(request, 'history.html', {'histories': histories})


def add_data_test(request):
    return render(request, 'DataEntry/add_data.html', {})